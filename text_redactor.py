import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText
from docx import Document
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
import ttkbootstrap as ttk
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO

def read_file(file_path):
    file_extension = file_path.split('.')[-1]
    if file_extension == 'txt':
        with open(file_path, 'r') as file:
            return file.read()
    elif file_extension == 'docx':
        doc = Document(file_path)
        return ' '.join([para.text for para in doc.paragraphs])
    elif file_extension == 'pdf':
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            return ' '.join([reader.pages[i].extract_text() for i in range(len(reader.pages))])

def redact_sensitive_information(text, patterns):
    for pattern, replacement in patterns.items():
        text = re.sub(pattern, replacement, text)
    return text

def open_file():
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt;*.docx;*.pdf"), ("All files", "*.*")])
    if not file_path:
        return
    text_widget.delete(1.0, tk.END)
    text = read_file(file_path)
    text_widget.insert(tk.END, text)
    analyze_text(text)

def save_file(file_format):
    redacted_text = text_widget.get(1.0, tk.END)
    file_path = filedialog.asksaveasfilename(defaultextension=f".{file_format}", filetypes=[(f"{file_format.upper()} files", f"*.{file_format}"), ("All files", "*.*")])
    if not file_path:
        return

    if file_format == 'txt':
        with open(file_path, 'w') as f:
            f.write(redacted_text)
    elif file_format == 'docx':
        doc = Document()
        doc.add_paragraph(redacted_text)
        doc.save(file_path)
    elif file_format == 'pdf':
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        text_object = can.beginText(40, 740)
        text_object.textLines(redacted_text.split('\n'))
        can.drawText(text_object)
        can.showPage()
        can.save()
        packet.seek(0)
        new_pdf = PdfReader(packet)
        output = PdfWriter()
        output.add_page(new_pdf.pages[0])
        with open(file_path, 'wb') as f:
            output.write(f)
    elif file_format == 'html':
        with open(file_path, 'w') as f:
            f.write(f"<html><body><pre>{redacted_text}</pre></body></html>")

    messagebox.showinfo("Success", f"{file_format.upper()} file saved.")

def clear_all():
    text_widget.delete(1.0, tk.END)
    recommendation_widget.delete(1.0, tk.END)
    for var in pattern_vars.values():
        var.set(1)

def toggle_all_redactions():
    state = all_redactions_var.get()
    for var in pattern_vars.values():
        var.set(state)

def analyze_text(text):
    recommendations = []
    for pattern, replacement in redaction_patterns.items():
        if re.search(pattern, text):
            recommendations.append(replacement)
    if recommendations:
        recommendation_text = "Recommended redactions:\n" + "\n".join(recommendations)
    else:
        recommendation_text = "No specific redactions recommended."
    recommendation_widget.delete(1.0, tk.END)
    recommendation_widget.insert(tk.END, recommendation_text)

def redact_text():
    text = text_widget.get(1.0, tk.END)
    selected_patterns = {pattern: replacement for pattern, replacement in redaction_patterns.items() if pattern_vars[pattern].get()}
    redacted_text = redact_sensitive_information(text, selected_patterns)
    text_widget.delete(1.0, tk.END)
    text_widget.insert(tk.END, redacted_text)

# Create the root window
root = ttk.Window(themename="darkly")
root.title("NeatLabs Rockin' Redactor")
root.geometry("1600x1600")

# Create text widget
text_widget = ScrolledText(root, wrap=tk.WORD, width=60, height=10)
text_widget.pack(expand=1, fill=tk.BOTH, padx=10, pady=10)

# Create buttons
button_frame = ttk.Frame(root)
button_frame.pack(fill=tk.X, padx=10, pady=10)

open_button = ttk.Button(button_frame, text="Open File", command=open_file)
open_button.pack(side=tk.LEFT, padx=5)

save_txt_button = ttk.Button(button_frame, text="Save as TXT", command=lambda: save_file('txt'))
save_txt_button.pack(side=tk.LEFT, padx=5)

save_docx_button = ttk.Button(button_frame, text="Save as DOCX", command=lambda: save_file('docx'))
save_docx_button.pack(side=tk.LEFT, padx=5)

save_pdf_button = ttk.Button(button_frame, text="Save as PDF", command=lambda: save_file('pdf'))
save_pdf_button.pack(side=tk.LEFT, padx=5)

save_html_button = ttk.Button(button_frame, text="Save as HTML", command=lambda: save_file('html'))
save_html_button.pack(side=tk.LEFT, padx=5)

clear_button = ttk.Button(button_frame, text="Clear", command=clear_all)
clear_button.pack(side=tk.LEFT, padx=5)

redact_button = ttk.Button(button_frame, text="Redact", command=redact_text)
redact_button.pack(side=tk.LEFT, padx=5)

# Create redaction pattern checkboxes
redaction_patterns = {
    r'\b\d{3}-\d{2}-\d{4}\b': '[REDACTED_SSN]',
    r'\b(?:\d{4}[-\s]?){3}\d{4}\b': '[REDACTED_CC]',
    r'\b(?:\d{1,3}\.){3}\d{1,3}\b': '[REDACTED_IP]',
    r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b': '[REDACTED_EMAIL]',
    r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b': '[REDACTED_IP]',
    r'\b[A-Za-z]\d[A-Za-z]\s?\d[A-Za-z]\d\b': '[REDACTED_POSTAL_CODE]',
    r'\b(?:\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}\b': '[REDACTED_PHONE]',
    r'\b[A-Z]{2}\d{2}[A-Z]{2}\d{4}\b': '[REDACTED_PASSPORT]',
    r'\b\d{3}-\d{4}-\d{4}-\d{4}\b': '[REDACTED_BANK_ACCOUNT]',
    r'\b[A-Z]{1,2}\d{1,6}[A-Z]?\b': '[REDACTED_LICENSE_PLATE]',
    r'\b(?:\d{1,3}\.){3}\d{1,3}:\d{1,5}\b': '[REDACTED_IP_PORT]',
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b': '[REDACTED_EMAIL]',
    r'\b(?:(?:https?|ftp):\/\/)?[\w/\-?=%.]+\.[\w/\-&?=%.]+\b': '[REDACTED_URL]'
}

pattern_vars = {}
checkbox_frame = ttk.Labelframe(root, text="Redaction Patterns", padding=10)
checkbox_frame.pack(fill=tk.X, padx=10, pady=10)

all_redactions_var = tk.BooleanVar(value=True)
all_redactions_checkbox = ttk.Checkbutton(checkbox_frame, text="Select All", variable=all_redactions_var, command=toggle_all_redactions)
all_redactions_checkbox.pack(anchor=tk.W)

for pattern, replacement in redaction_patterns.items():
    var = tk.BooleanVar(value=True)
    pattern_vars[pattern] = var
    checkbox = ttk.Checkbutton(checkbox_frame, text=replacement, variable=var)
    checkbox.pack(anchor=tk.W)

# Create recommendation widget
recommendation_frame = ttk.Labelframe(root, text="Redaction Recommendations", padding=10)
recommendation_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

recommendation_widget = ScrolledText(recommendation_frame, wrap=tk.WORD, width=60, height=10)
recommendation_widget.pack(fill=tk.BOTH, expand=True)

root.mainloop()